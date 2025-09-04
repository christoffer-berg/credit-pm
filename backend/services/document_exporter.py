from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from typing import Dict, List, Any
from datetime import datetime
import os

def export_to_word(case: Dict, company: Dict, sections: List[Dict], output_path: str) -> str:
    """
    Export PM case to Word document using python-docx.
    """
    try:
        # Create document
        doc = Document()
        
        # Set up styles
        _setup_word_styles(doc)
        
        # Add title
        title = doc.add_heading(case.get('title', 'Credit Memo'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company information
        _add_company_info_word(doc, company)
        
        # Add case metadata
        _add_case_metadata_word(doc, case)
        
        # Add sections
        section_order = [
            'purpose', 'business_description', 'market_analysis',
            'financial_analysis', 'credit_analysis', 'credit_proposal'
        ]
        
        for section_type in section_order:
            section = next((s for s in sections if s.get('section_type') == section_type), None)
            if section:
                _add_section_word(doc, section)
        
        # Add footer with metadata
        _add_footer_word(doc, case)
        
        # Save document
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        raise Exception(f"Error creating Word document: {str(e)}")

def export_to_pdf(case: Dict, company: Dict, sections: List[Dict], output_path: str) -> str:
    """
    Export PM case to PDF using reportlab.
    """
    try:
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for elements
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceBefore=6,
            spaceAfter=6,
            leading=12
        )
        
        # Add title
        story.append(Paragraph(case.get('title', 'Credit Memo'), title_style))
        story.append(Spacer(1, 12))
        
        # Add company information
        _add_company_info_pdf(story, company, styles)
        
        # Add case metadata
        _add_case_metadata_pdf(story, case, styles)
        
        # Add sections
        section_order = [
            'purpose', 'business_description', 'market_analysis',
            'financial_analysis', 'credit_analysis', 'credit_proposal'
        ]
        
        for section_type in section_order:
            section = next((s for s in sections if s.get('section_type') == section_type), None)
            if section:
                _add_section_pdf(story, section, heading_style, body_style)
        
        # Add footer
        _add_footer_pdf(story, case, styles)
        
        # Build PDF
        doc.build(story)
        return output_path
        
    except Exception as e:
        raise Exception(f"Error creating PDF document: {str(e)}")

def _setup_word_styles(doc: Document):
    """Set up custom styles for Word document."""
    styles = doc.styles
    
    # Create custom heading style
    try:
        heading_style = styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.base_style = styles['Heading 2']
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
    except:
        pass  # Style might already exist

def _add_company_info_word(doc: Document, company: Dict):
    """Add company information section to Word document."""
    if not company:
        return
    
    doc.add_heading('Company Information', level=2)
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    info_items = [
        ('Company Name', company.get('name', 'N/A')),
        ('Organization Number', company.get('organization_number', 'N/A')),
        ('Industry Code', company.get('industry_code', 'N/A')),
        ('Business Description', company.get('business_description', 'N/A'))
    ]
    
    for label, value in info_items:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)
    
    doc.add_paragraph()

def _add_case_metadata_word(doc: Document, case: Dict):
    """Add case metadata to Word document."""
    doc.add_heading('Case Information', level=2)
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    metadata_items = [
        ('Status', case.get('status', 'N/A').title()),
        ('Version', str(case.get('version', 'N/A'))),
        ('Created', case.get('created_at', 'N/A')[:10] if case.get('created_at') else 'N/A'),
        ('Updated', case.get('updated_at', 'N/A')[:10] if case.get('updated_at') else 'N/A')
    ]
    
    for label, value in metadata_items:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    
    doc.add_paragraph()

def _add_section_word(doc: Document, section: Dict):
    """Add a section to Word document."""
    title = section.get('title', section.get('section_type', 'Section')).title()
    doc.add_heading(title, level=2)
    
    # Use user content if available, otherwise AI content
    content = section.get('user_content') or section.get('ai_content', 'No content available')
    
    paragraphs = content.split('\n\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    doc.add_paragraph()

def _add_footer_word(doc: Document, case: Dict):
    """Add footer with document metadata."""
    doc.add_page_break()
    doc.add_heading('Document Information', level=2)
    
    footer_info = [
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Case Version: {case.get('version', 'N/A')}",
        f"Data Sources: Bolagsverket API, Internal Analysis",
        f"AI Model Version: GPT-4",
        "",
        "This document was generated automatically by the Credit PM Generator system."
    ]
    
    for info in footer_info:
        p = doc.add_paragraph(info)
        if info.startswith("This document"):
            p.italic = True

def _add_company_info_pdf(story: List, company: Dict, styles):
    """Add company information section to PDF."""
    if not company:
        return
    
    story.append(Paragraph('<b>Company Information</b>', styles['Heading2']))
    story.append(Spacer(1, 12))
    
    data = [
        ['Company Name', company.get('name', 'N/A')],
        ['Organization Number', company.get('organization_number', 'N/A')],
        ['Industry Code', company.get('industry_code', 'N/A')],
        ['Business Description', company.get('business_description', 'N/A')]
    ]
    
    table = Table(data, colWidths=[2*inch, 4*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(table)
    story.append(Spacer(1, 20))

def _add_case_metadata_pdf(story: List, case: Dict, styles):
    """Add case metadata to PDF."""
    story.append(Paragraph('<b>Case Information</b>', styles['Heading2']))
    story.append(Spacer(1, 12))
    
    data = [
        ['Status', case.get('status', 'N/A').title()],
        ['Version', str(case.get('version', 'N/A'))],
        ['Created', case.get('created_at', 'N/A')[:10] if case.get('created_at') else 'N/A'],
        ['Updated', case.get('updated_at', 'N/A')[:10] if case.get('updated_at') else 'N/A']
    ]
    
    table = Table(data, colWidths=[2*inch, 4*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(table)
    story.append(Spacer(1, 20))

def _add_section_pdf(story: List, section: Dict, heading_style, body_style):
    """Add a section to PDF."""
    title = section.get('title', section.get('section_type', 'Section')).title()
    story.append(Paragraph(f'<b>{title}</b>', heading_style))
    
    # Use user content if available, otherwise AI content
    content = section.get('user_content') or section.get('ai_content', 'No content available')
    
    paragraphs = content.split('\n\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            # Escape HTML characters in content
            escaped_content = paragraph.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(escaped_content, body_style))
    
    story.append(Spacer(1, 12))

def _add_footer_pdf(story: List, case: Dict, styles):
    """Add footer with document metadata."""
    story.append(Spacer(1, 30))
    story.append(Paragraph('<b>Document Information</b>', styles['Heading2']))
    story.append(Spacer(1, 12))
    
    footer_info = [
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Case Version: {case.get('version', 'N/A')}",
        f"Data Sources: Bolagsverket API, Internal Analysis",
        f"AI Model Version: GPT-4",
        "",
        "<i>This document was generated automatically by the Credit PM Generator system.</i>"
    ]
    
    for info in footer_info:
        if info:
            story.append(Paragraph(info, styles['Normal']))
            story.append(Spacer(1, 6))